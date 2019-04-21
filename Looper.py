# -*- coding: utf-8 -*-
"""
Created on Fri Apr 27 17:54:56 2018

@author: Miha
"""

import os
from docx import Document
import io



def doc_writer(strings, lokacija_ind):
    obcina_ind = improved_search(strings, 'Občina:', 23)
    obcina = data_nicefier(strings[obcina_ind + 1]).upper()
    path = 'TEKSTI/' + obcina + '.docx'
    data = data_stringyfier(strings, lokacija_ind)
    if os.path.exists(path):
        document = Document(path)
        paragraph = document.add_paragraph()
        paragraph.add_run(data[0]).bold = True
        paragraph.add_run(data[1])
        paragraph.add_run('Datacija enote: ')
        paragraph.add_run(data[2]).bold = True
        paragraph.add_run('Lokacija: ').bold = True
        paragraph.add_run(data[3])
        document.save(path)
        return None
    else:
        helper1 = "OBČINA {0}- OBČINA {0}- OBČINA {0}- OBČINA {0}\nREGISTER NEPREMIČNINE KULTURNE DEDIŠČINE \nREPUBLIKA SLOVENIJA \nMINISTRSTVO ZA KULTURO\nPOGOJI UPORABE\n\n".format((obcina))
        helper2 = "Pri uporabi podatkov iz registra mora vsak uporabnik register navesti kot vir podatkov (Zakon o varstvu kulturne dediščine, Uradni list RS, št. 16/2008).\n\n"
        helper3 = "SEZNAM NEPREMIČNE KULTURNE DEDIŠČINE – OBČINA {0}\n\n\n\n".format((obcina))
        document = Document()
        paragraph = document.add_paragraph()
        paragraph.add_run(helper1).bold = True
        paragraph.add_run(helper2)
        paragraph.add_run(helper3).bold = True
        paragraph.add_run(data[0]).bold = True
        paragraph.add_run(data[1])
        paragraph.add_run('Datacija enote: ')
        paragraph.add_run(data[2]).bold = True
        paragraph.add_run('Lokacija: ').bold = True
        paragraph.add_run(data[3])
        document.save(path)
        return None





def main_loop():
    itera = [folder for folder in os.listdir('TEKSTI')]
    for folder in itera[352: ]:
        one_loop(folder)

def one_loop(folder):
    path = 'TEKSTI/' + folder
    imena = [file for file in os.listdir(path)]
    for znam in imena:
        read2write(znam, folder)
        

def read2write(znam, folder):
    with open('TEKSTI/' + folder + '/' + znam, 'r', encoding = 'utf-8') as file:
        lines = [line for line in file.readlines() if line != '\n']
        path = 'TEKSTI/' + folder + '.docx'
        ime = znam[:len(znam) - 4] + '\n' + '\n'
        if len(lines) == 4:
            tekst = lines[1] + '\n'
            datacija = lines[2] +'\n'
            lokacija = lines[3] + '\n'
        if len(lines) == 3:
            celi_tekst = lines[0]
            tekst = celi_tekst[len(ime) - 2:] + '\n'
            datacija = lines[1] + '\n'
            lokacija = lines[2] + '\n'
                
        if os.path.exists(path):
            document = Document(path)
            paragraph = document.add_paragraph()
            paragraph.add_run(ime).bold = True
            paragraph.add_run(tekst)
            paragraph.add_run('Datacija enote: ')
            paragraph.add_run(datacija).bold = True
            paragraph.add_run('Lokacija: ').bold = True
            paragraph.add_run(lokacija)
            document.save(path)
            return None
        else:
            helper1 = "OBČINA {0}- OBČINA {0}- OBČINA {0}- OBČINA {0}\nREGISTER NEPREMIČNINE KULTURNE DEDIŠČINE \nREPUBLIKA SLOVENIJA \nMINISTRSTVO ZA KULTURO\nPOGOJI UPORABE\n\n".format((folder))
            helper2 = "Pri uporabi podatkov iz registra mora vsak uporabnik register navesti kot vir podatkov (Zakon o varstvu kulturne dediščine, Uradni list RS, št. 16/2008).\n\n"
            helper3 = "SEZNAM NEPREMIČNE KULTURNE DEDIŠČINE – OBČINA {0}\n\n\n\n".format((folder))
            document = Document()
            paragraph = document.add_paragraph()
            paragraph.add_run(helper1).bold = True
            paragraph.add_run(helper2)
            paragraph.add_run(helper3).bold = True
            paragraph.add_run(ime).bold = True
            paragraph.add_run(tekst)
            paragraph.add_run('Datacija enote: ')
            paragraph.add_run(datacija).bold = True
            paragraph.add_run('Lokacija: ').bold = True
            paragraph.add_run(lokacija)
            document.save(path)
            return None




#quicksort
def partition(alist, left, right):
    """Partitions a list in order to quicksort it."""
    v = alist[left]
    m = left
    for i in range(left + 1, right):
        if alist[i] < v:
            m += 1
            (alist[i], alist[m]) = (alist[m], alist[i])
    (alist[left], alist[m]) = (alist[m], alist[left])
    return m
    

def quicksort(alist):
    def quicksort_rec(alist, left, right):
        if left != right:
            m = partition(alist, left, right)
            quicksort_rec(alist, left, m)
            quicksort_rec(alist, m + 1, right)
        
    quicksort_rec(alist, 0, len(alist))
    return alist