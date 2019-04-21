# -*- coding: utf-8 -*-
"""
Created on Fri Jan 12 19:28:36 2018

@author: Miha
"""
import os.path
from docx import Document
import io



def doc_writer(strings, lokacija_ind):
    obcina_ind = improved_search(strings, 'Občina:', 23)
    obcina = data_nicefier(strings[obcina_ind + 1]).upper()
    path = 'TEKSTI/' + obcina + '.docx'
    data = data_stringyfier(strings, lokacija_ind)
    if os.path.exists(path):
        document = Document(path)
        paragraph = document.add_paragraph()
        paragraph.add_run(data[0]).bold = True
        paragraph.add_run(data[1])
        paragraph.add_run('Datacija enote: ')
        paragraph.add_run(data[2]).bold = True
        paragraph.add_run('Lokacija: ').bold = True
        paragraph.add_run(data[3])
        document.save(path)
        return None
    else:
        helper1 = "OBČINA {0}- OBČINA {0}- OBČINA {0}- OBČINA {0}\nREGISTER NEPREMIČNINE KULTURNE DEDIŠČINE \nREPUBLIKA SLOVENIJA \nMINISTRSTVO ZA KULTURO\nPOGOJI UPORABE\n\n".format((obcina))
        helper2 = "Pri uporabi podatkov iz registra mora vsak uporabnik register navesti kot vir podatkov (Zakon o varstvu kulturne dediščine, Uradni list RS, št. 16/2008).\n\n"
        helper3 = "SEZNAM NEPREMIČNE KULTURNE DEDIŠČINE – OBČINA {0}\n\n\n\n".format((obcina))
        document = Document()
        paragraph = document.add_paragraph()
        paragraph.add_run(helper1).bold = True
        paragraph.add_run(helper2)
        paragraph.add_run(helper3).bold = True
        paragraph.add_run(data[0]).bold = True
        paragraph.add_run(data[1])
        paragraph.add_run('Datacija enote: ')
        paragraph.add_run(data[2]).bold = True
        paragraph.add_run('Lokacija: ').bold = True
        paragraph.add_run(data[3])
        document.save(path)
        return None


def file_writer(strings, lokacija_ind):
    """Because word is not sufficient for my needs, this function
    takes in a list of strings"""
    obcina_ind = improved_search(strings, 'Občina:', 23)
    obcina = data_nicefier(strings[obcina_ind + 1]).upper()
    data = data_stringyfier(strings, lokacija_ind)
    ime, tekst, datacija, lokacija = data
    ime = ime.strip()
    ime = eliminator(ime)
    path = 'TEKSTI/' + obcina + '/' + ime.strip() + '.txt'
    if not os.path.exists('TEKSTI/' + obcina):
        os.makedirs('TEKSTI/' + obcina)
    with open(path, 'w', encoding="utf-8") as file:
        file.write(ime + tekst + datacija + lokacija)
        
    
def eliminator(string):
    """ELIMINAATTATATATATES"""
    for i in range(len(string)):
        if string[i] == '/' or string[i] == "\\":
            string = string[:i] + '_' + string[i+1:]
    return string
    
    
def data_stringyfier(strings, lokacija_ind):
    """Takes in a list of strings and returns a single tuple with the relevant data"""
    ime = data_nicefier(strings[5]) + '\n' + '\n'
    tekst = strings[16] + '\n' + '\n'
    datacija =data_nicefier(strings[18]) + '\n' + '\n'
    
    if selector(strings, lokacija_ind):
        lokacija = data_nicefier(strings[lokacija_ind + 1]) + '\n' + '\n'
    
    else:
        lokacija = data_nicefier(strings[lokacija_ind - 3]) + '\n' + '\n'
 
    string = (ime, tekst, datacija, lokacija)
    return string


def selector(strings, lokacija_ind):
    """Takes in a list of strings and determines whether they are useful or not."""
    return not ((len(strings) < 5) or (strings[lokacija_ind + 1] == 'PRISTOJNOSTI'))


def improved_search(strings, searched, start):
    """Takes in three arguments: a list of 'strings', a string 'searched', and
    an integer start. The functions compares each element of the list to the
    string 'searched' and returns the index at which it finds the match."""
    for i in range(start, len(strings)):
        if strings[i] == searched:
            return i
        
def data_nicefier(string):
    """Takes in a 'string' and makes it nicer by stripping away all the unecessary parts
    and properly formatting it."""
    return ' '.join(string.split())